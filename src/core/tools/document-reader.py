from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


def normalize_text(value: object) -> str:
    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def non_empty_lines(values: list[str]) -> list[str]:
    return [value for value in (normalize_text(item) for item in values) if value]


def extract_pdf(file_path: Path) -> dict[str, object]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    blocks: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if text:
            blocks.append(f"[Page {index}]\n{text}")

    return {
        "format": "pdf",
        "content": "\n\n".join(blocks) or "No extractable text was found in this PDF.",
        "metadata": {
            "pages": len(reader.pages),
        },
    }


def extract_docx(file_path: Path) -> dict[str, object]:
    from docx import Document

    document = Document(str(file_path))
    blocks = non_empty_lines([paragraph.text for paragraph in document.paragraphs])

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells).rstrip())

    return {
        "format": "docx",
        "content": "\n".join(blocks) or "No extractable text was found in this document.",
        "metadata": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
        },
    }


def normalize_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return normalize_text(value)


def extract_xlsx(file_path: Path) -> dict[str, object]:
    from openpyxl import load_workbook

    workbook = load_workbook(str(file_path), data_only=True, read_only=True)
    sheet_blocks: list[str] = []
    sheet_names: list[str] = []

    for sheet in workbook.worksheets:
        sheet_names.append(sheet.title)
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [normalize_cell(cell) for cell in row]
            if any(values):
                rows.append("\t".join(values).rstrip())

        if rows:
            sheet_blocks.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))

    workbook.close()

    return {
        "format": "xlsx",
        "content": "\n\n".join(sheet_blocks) or "No readable cell data was found in this workbook.",
        "metadata": {
            "sheetCount": len(sheet_names),
            "sheets": sheet_names,
        },
    }


def natural_order(name: str) -> list[object]:
    return [int(part) if part.isdigit() else part for part in re.split(r"(\d+)", name)]


def extract_pptx(file_path: Path) -> dict[str, object]:
    namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    slide_blocks: list[str] = []

    with zipfile.ZipFile(file_path) as archive:
        slide_paths = sorted(
            [
                item
                for item in archive.namelist()
                if item.startswith("ppt/slides/slide") and item.endswith(".xml")
            ],
            key=natural_order,
        )

        for index, slide_path in enumerate(slide_paths, start=1):
            root = ET.fromstring(archive.read(slide_path))
            texts = non_empty_lines(
                [node.text or "" for node in root.findall(".//a:t", namespace)]
            )
            if texts:
                slide_blocks.append(f"[Slide {index}]\n" + "\n".join(texts))

    return {
        "format": "pptx",
        "content": "\n\n".join(slide_blocks) or "No extractable text was found in this presentation.",
        "metadata": {
            "slides": len(slide_blocks),
        },
    }


def extract_open_document(file_path: Path) -> dict[str, object]:
    namespace = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
    }

    with zipfile.ZipFile(file_path) as archive:
        root = ET.fromstring(archive.read("content.xml"))

    paragraphs = non_empty_lines(
        [node.text or "" for node in root.findall(".//text:p", namespace)]
    )
    headings = non_empty_lines(
        [node.text or "" for node in root.findall(".//text:h", namespace)]
    )
    rows = []
    for row in root.findall(".//table:table-row", namespace):
        cells = non_empty_lines(
            [cell.text or "" for cell in row.findall(".//text:p", namespace)]
        )
        if cells:
            rows.append("\t".join(cells))

    blocks = headings + paragraphs + rows
    extension = file_path.suffix.lower().lstrip(".")

    return {
        "format": extension,
        "content": "\n".join(blocks) or "No extractable text was found in this OpenDocument file.",
        "metadata": {
            "paragraphs": len(paragraphs),
            "rows": len(rows),
        },
    }


def extract(file_path: Path) -> dict[str, object]:
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf(file_path)
    if extension == ".docx":
        return extract_docx(file_path)
    if extension in {".xlsx", ".xlsm"}:
        return extract_xlsx(file_path)
    if extension == ".pptx":
        return extract_pptx(file_path)
    if extension in {".odt", ".ods", ".odp"}:
        return extract_open_document(file_path)

    raise RuntimeError(f"Unsupported structured document format: {extension or 'unknown'}")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: document-reader.py <file>", file=sys.stderr)
        return 1

    file_path = Path(sys.argv[1]).expanduser().resolve()

    try:
        payload = extract(file_path)
    except Exception as error:  # noqa: BLE001
        print(str(error), file=sys.stderr)
        return 1

    sys.stdout.buffer.write(json.dumps(payload, ensure_ascii=False).encode("utf-8"))
    sys.stdout.buffer.write(b"\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
